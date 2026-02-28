#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同文档解析脚本 - Windows 兼容版
支持格式：PDF, DOCX
输出格式：Markdown
"""

import sys
import argparse
from pathlib import Path

def parse_docx(input_path: Path, output_path: Path) -> str:
    """解析 DOCX 文件"""
    import docx
    
    doc = docx.Document(input_path)
    content = []
    
    content.append(f"# {input_path.name}\n\n")
    content.append("---\n\n")
    
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
            content.append("\n\n")
    
    # 提取表格
    tables = doc.tables
    if tables:
        content.append("## 合同表格\n\n")
        for t_idx, table in enumerate(tables, 1):
            content.append(f"### 表格 {t_idx}\n\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append("| " + " | ".join(row_text) + " |\n")
            content.append("\n")
    
    return "\n".join(content)


def parse_pdf(input_path: Path, output_path: Path) -> str:
    """解析 PDF 文件"""
    from PyPDF2 import PdfReader
    
    reader = PdfReader(input_path)
    content = []
    
    content.append(f"# {input_path.name}\n")
    content.append(f"**页数**: {len(reader.pages)}\n\n")
    content.append("---\n\n")
    
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            content.append(f"## 第 {i} 页\n\n")
            content.append(text)
            content.append("\n\n---\n\n")
    
    return "\n".join(content)


def main():
    # 使用 Windows 原生 API 处理路径
    if len(sys.argv) < 2:
        print("用法：python parse_doc_simple.py <输入文件> [输出目录]")
        sys.exit(1)
    
    # 通过命令行参数获取文件路径（避免编码问题）
    input_arg = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "text_cache"
    
    # 查找 contracts 目录中的文件
    contracts_dir = Path(__file__).parent / "contracts"
    output_dir_path = Path(__file__).parent / output_dir
    output_dir_path.mkdir(parents=True, exist_ok=True)
    
    # 匹配文件
    input_file = None
    for f in contracts_dir.glob("*.docx"):
        if input_arg in str(f) or input_arg in f.name:
            input_file = f
            break
    
    if not input_file:
        for f in contracts_dir.glob("*.pdf"):
            if input_arg in str(f) or input_arg in f.name:
                input_file = f
                break
    
    if not input_file:
        print(f"[ERROR] 找不到文件：{input_arg}")
        print("可用文件:")
        for f in contracts_dir.glob("*"):
            print(f"  - {f.name}")
        sys.exit(1)
    
    print(f"[INFO] 正在解析：{input_file.name}")
    
    try:
        suffix = input_file.suffix.lower()
        output_file = output_dir_path / f"{input_file.stem}.md"
        
        if suffix == ".docx":
            content = parse_docx(input_file, output_file)
        elif suffix == ".pdf":
            content = parse_pdf(input_file, output_file)
        else:
            print(f"[ERROR] 不支持的格式：{suffix}")
            sys.exit(1)
        
        # 写入文件
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)
        
        print(f"[OK] 解析完成：{output_file.name}")
        print(f"[INFO] 文件位置：{output_dir}\\{output_file.name}")
        
    except Exception as e:
        print(f"[ERROR] 解析失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
