#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同文档解析脚本 - Windows 兼容版
使用 os.listdir 避免中文路径编码问题
"""

import os
import sys
import argparse

def parse_docx(input_path: str, output_path: str) -> str:
    """解析 DOCX 文件"""
    import docx
    
    doc = docx.Document(input_path)
    content = []
    
    filename = os.path.basename(input_path)
    content.append(f"# {filename}\n\n")
    content.append("---\n\n")
    
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
            content.append("\n\n")
    
    # 提取表格
    tables = doc.tables
    if tables:
        content.append("## 合同表格\n\n")
        for t_idx, table in enumerate(tables, 1):
            content.append(f"### 表格 {t_idx}\n\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append("| " + " | ".join(row_text) + " |\n")
            content.append("\n")
    
    return "\n".join(content)


def parse_pdf(input_path: str, output_path: str) -> str:
    """解析 PDF 文件"""
    from PyPDF2 import PdfReader
    
    reader = PdfReader(input_path)
    content = []
    
    filename = os.path.basename(input_path)
    content.append(f"# {filename}\n")
    content.append(f"**页数**: {len(reader.pages)}\n\n")
    content.append("---\n\n")
    
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            content.append(f"## 第 {i} 页\n\n")
            content.append(text)
            content.append("\n\n---\n\n")
    
    return "\n".join(content)


def find_file(contracts_dir: str, search_term: str) -> str:
    """在 contracts 目录中查找文件"""
    files = os.listdir(contracts_dir)
    
    # 优先匹配 docx
    for f in files:
        if f.endswith('.docx') and search_term in f:
            return os.path.join(contracts_dir, f)
    
    # 其次匹配 pdf
    for f in files:
        if f.endswith('.pdf') and search_term in f:
            return os.path.join(contracts_dir, f)
    
    # 如果搜索词没找到，返回第一个 docx 或 pdf
    for f in files:
        if f.endswith('.docx'):
            return os.path.join(contracts_dir, f)
    
    for f in files:
        if f.endswith('.pdf'):
            return os.path.join(contracts_dir, f)
    
    return None


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(script_dir)
    contracts_dir = os.path.join(parent_dir, "contracts")
    output_dir = os.path.join(parent_dir, "text_cache")
    
    os.makedirs(output_dir, exist_ok=True)
    
    if len(sys.argv) < 2:
        print("[INFO] 用法：python parse_doc_win.py <文件关键词>")
        print("[INFO] 示例：python parse_doc_win.py 合肥晶合")
        print("\n[INFO] 可用文件:")
        for f in os.listdir(contracts_dir):
            print(f"  - {f}")
        sys.exit(0)
    
    search_term = sys.argv[1]
    input_file = find_file(contracts_dir, search_term)
    
    if not input_file:
        print(f"[ERROR] 找不到文件，搜索词：{search_term}")
        print("\n可用文件:")
        for f in os.listdir(contracts_dir):
            print(f"  - {f}")
        sys.exit(1)
    
    print(f"[INFO] 正在解析：{os.path.basename(input_file)}")
    
    try:
        suffix = os.path.basename(input_file).lower()
        if suffix.endswith('.docx'):
            suffix = '.docx'
        elif suffix.endswith('.pdf'):
            suffix = '.pdf'
        else:
            suffix = os.path.splitext(input_file)[1].lower()
        
        output_filename = os.path.splitext(os.path.basename(input_file))[0] + ".md"
        output_file = os.path.join(output_dir, output_filename)
        
        if suffix == ".docx":
            content = parse_docx(input_file, output_file)
        elif suffix == ".pdf":
            content = parse_pdf(input_file, output_file)
        else:
            print(f"[ERROR] 不支持的格式：{suffix}")
            sys.exit(1)
        
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)
        
        print(f"[OK] 解析完成！")
        print(f"[INFO] 输出文件：text_cache\\{output_filename}")
        print(f"[INFO] 完整路径：{output_file}")
        
    except Exception as e:
        print(f"[ERROR] 解析失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
