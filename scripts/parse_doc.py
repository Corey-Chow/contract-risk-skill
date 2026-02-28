#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同文档解析脚本
支持格式：PDF, DOCX, DOC
输出格式：Markdown
"""

import os
import sys
import argparse
import locale
from pathlib import Path

# 设置控制台编码为 UTF-8
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# 依赖检查
def check_dependencies():
    missing = []
    try:
        import PyPDF2
    except ImportError:
        missing.append("PyPDF2")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    
    if missing:
        print(f"缺少依赖库：{missing}")
        print(f"请运行：pip install {' '.join(missing)}")
        return False
    return True


def parse_pdf(input_path: str, output_path: str) -> str:
    """解析 PDF 文件"""
    from PyPDF2 import PdfReader
    
    reader = PdfReader(input_path)
    content = []
    
    content.append(f"# {Path(input_path).name}\n")
    content.append(f"**页数**: {len(reader.pages)}\n\n")
    content.append("---\n\n")
    
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            content.append(f"## 第 {i} 页\n\n")
            content.append(text)
            content.append("\n\n---\n\n")
    
    return "\n".join(content)


def parse_docx(input_path: str, output_path: str) -> str:
    """解析 DOCX 文件"""
    import docx
    
    doc = docx.Document(input_path)
    content = []
    
    content.append(f"# {Path(input_path).name}\n\n")
    content.append("---\n\n")
    
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
            content.append("\n\n")
    
    # 提取表格
    tables = doc.tables
    if tables:
        content.append("## 合同表格\n\n")
        for t_idx, table in enumerate(tables, 1):
            content.append(f"### 表格 {t_idx}\n\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append("| " + " | ".join(row_text) + " |\n")
            content.append("\n")
    
    return "\n".join(content)


def parse_document(input_path: str, output_dir: str) -> str:
    """根据文件类型选择解析器"""
    input_p = Path(input_path)
    output_p = Path(output_dir)
    output_p.mkdir(parents=True, exist_ok=True)
    
    suffix = input_p.suffix.lower()
    output_filename = f"{input_p.stem}.md"
    output_path = output_p / output_filename
    
    print(f"正在解析：{input_path}")
    
    if suffix == ".pdf":
        content = parse_pdf(str(input_path), str(output_path))
    elif suffix in [".docx", ".doc"]:
        content = parse_docx(str(input_path), str(output_path))
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")
    
    # 写入输出文件
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"解析完成：{output_path}")
    return str(output_path)


def main():
    parser = argparse.ArgumentParser(description="合同文档解析工具")
    parser.add_argument("input", help="输入文件路径 (PDF/DOCX)")
    parser.add_argument("-o", "--output", default="text_cache", 
                        help="输出目录 (默认：text_cache)")
    
    args = parser.parse_args()
    
    if not check_dependencies():
        sys.exit(1)
    
    try:
        output_file = parse_document(args.input, args.output)
        print(f"\n[OK] 解析成功，输出文件：{output_file}")
    except Exception as e:
        print(f"\n[ERROR] 解析失败：{e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
