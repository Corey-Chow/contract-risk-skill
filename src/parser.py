#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同文档解析模块
支持格式：PDF, DOCX
输出：Markdown 文本
"""

import os
from pathlib import Path


def parse_docx(file_path: Path) -> str:
    """解析 DOCX 文件"""
    import docx
    
    doc = docx.Document(str(file_path))
    content = []
    
    content.append(f"# {file_path.name}\n\n")
    content.append("---\n\n")
    
    # 提取段落
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
            content.append("\n\n")
    
    # 提取表格
    if doc.tables:
        content.append("## 合同表格\n\n")
        for t_idx, table in enumerate(doc.tables, 1):
            content.append(f"### 表格 {t_idx}\n\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    content.append("| " + " | ".join(row_text) + " |\n")
            content.append("\n")
    
    return "\n".join(content)


def parse_pdf(file_path: Path) -> str:
    """解析 PDF 文件"""
    from PyPDF2 import PdfReader
    
    reader = PdfReader(str(file_path))
    content = []
    
    content.append(f"# {file_path.name}\n")
    content.append(f"**页数**: {len(reader.pages)}\n\n")
    content.append("---\n\n")
    
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            content.append(f"## 第 {i} 页\n\n")
            content.append(text)
            content.append("\n\n---\n\n")
    
    return "\n".join(content)


def parse_document(file_path: str) -> str:
    """根据文件类型选择解析器"""
    path = Path(file_path)
    
    if not path.exists():
        # 尝试在 contracts 目录查找
        contracts_dir = Path(__file__).parent.parent / "contracts"
        if contracts_dir.exists():
            for f in contracts_dir.iterdir():
                if file_path in str(f) or file_path in f.name:
                    path = f
                    break
    
    if not path.exists():
        raise FileNotFoundError(f"找不到文件：{file_path}")
    
    suffix = path.suffix.lower()
    
    if suffix == ".docx":
        return parse_docx(path)
    elif suffix == ".pdf":
        return parse_pdf(path)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")
